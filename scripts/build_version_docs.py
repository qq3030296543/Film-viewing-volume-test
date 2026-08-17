from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "Docs"

COLORS = {
    "navy": "0B2545",
    "blue": "2E74B5",
    "dark_blue": "1F4D78",
    "muted": "667085",
    "gold": "A97835",
    "body": "20242A",
    "light_fill": "E8EEF5",
    "border": "C8D2DE",
    "white": "FFFFFF",
}


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_run_font(run, size=None, bold=None, italic=None, color=None, latin="Calibri", east_asia="Microsoft YaHei"):
    run.font.name = latin
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), latin)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), latin)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = rgb(color)


def set_style_font(style, latin="Calibri", east_asia="Microsoft YaHei", size=11, color="20242A"):
    style.font.name = latin
    style.font.size = Pt(size)
    style.font.color.rgb = rgb(color)
    rpr = style.element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), latin)
    rpr.rFonts.set(qn("w:hAnsi"), latin)
    rpr.rFonts.set(qn("w:eastAsia"), east_asia)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa=120):
    total = sum(widths_dxa)
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def mark_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def add_page_field(paragraph):
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=COLORS["muted"])
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    field_sep = OxmlElement("w:fldChar")
    field_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    fld_run = paragraph.add_run()
    set_run_font(fld_run, size=9, color=COLORS["muted"])
    fld_run._r.extend([field_begin, instr, field_sep, text, field_end])
    tail = paragraph.add_run(" 页")
    set_run_font(tail, size=9, color=COLORS["muted"])


def configure_document(doc: Document, short_title: str):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    set_style_font(normal, size=11, color=COLORS["body"])
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, COLORS["blue"], 18, 10),
        "Heading 2": (13, COLORS["blue"], 14, 7),
        "Heading 3": (12, COLORS["dark_blue"], 10, 5),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[name]
        set_style_font(style, size=size, color=color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True

    header = section.header
    header_p = header.paragraphs[0]
    header_p.paragraph_format.space_after = Pt(0)
    header_p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    left = header_p.add_run("光影鉴赏局 · 版本更新记录")
    set_run_font(left, size=8.5, bold=True, color=COLORS["muted"])
    right = header_p.add_run(f"\t{short_title}")
    set_run_font(right, size=8.5, color=COLORS["muted"])

    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_p.paragraph_format.space_before = Pt(0)
    add_page_field(footer_p)

    props = doc.core_properties
    props.author = "光影鉴赏局项目组"
    props.subject = "电影阅历测试网站版本更新文档"
    props.keywords = "电影测试, TMDB, React, TypeScript, Vite, Netlify"


def add_title_block(doc: Document, version: str, title: str, summary: str, status="已完成"):
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(12)
    kicker.paragraph_format.space_after = Pt(5)
    run = kicker.add_run(f"CINE MEMORY BUREAU  /  {version}")
    set_run_font(run, size=9, bold=True, color=COLORS["gold"])

    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(6)
    title_p.paragraph_format.keep_with_next = True
    run = title_p.add_run(title)
    set_run_font(run, size=27, bold=True, color=COLORS["navy"])

    summary_p = doc.add_paragraph()
    summary_p.paragraph_format.space_before = Pt(0)
    summary_p.paragraph_format.space_after = Pt(12)
    summary_p.paragraph_format.line_spacing = 1.2
    run = summary_p.add_run(summary)
    set_run_font(run, size=11.5, color=COLORS["muted"])

    meta = doc.add_paragraph()
    meta.paragraph_format.space_before = Pt(0)
    meta.paragraph_format.space_after = Pt(14)
    for idx, text in enumerate((f"版本：{version}", "文档类型：功能更新说明", f"状态：{status}")):
        if idx:
            sep = meta.add_run("   |   ")
            set_run_font(sep, size=9, color=COLORS["border"])
        run = meta.add_run(text)
        set_run_font(run, size=9, bold=idx == 0, color=COLORS["muted"])


def create_numbering(doc: Document, bullet=True) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(el.get(qn("w:abstractNumId"))) for el in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(el.get(qn("w:numId"))) for el in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if bullet else "decimal")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if bullet else "%1.")
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.extend([tabs, ind, spacing])
    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    fonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r_pr.append(fonts)
    lvl.extend([start, num_fmt, lvl_text, lvl_jc, p_pr, r_pr])
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abs_id = OxmlElement("w:abstractNumId")
    abs_id.set(qn("w:val"), str(abstract_id))
    num.append(abs_id)
    numbering.append(num)
    return num_id


def add_list_item(doc: Document, text: str, num_id: int):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.keep_together = True
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])
    p_pr.append(num_pr)
    if "：" in text:
        label, detail = text.split("：", 1)
        label_run = p.add_run(f"{label}：")
        set_run_font(label_run, size=11, bold=True, color=COLORS["body"])
        detail_run = p.add_run(detail)
        set_run_font(detail_run, size=11, color=COLORS["body"])
    else:
        run = p.add_run(text)
        set_run_font(run, size=11, color=COLORS["body"])
    return p


def add_bullets(doc: Document, items: Iterable[str], num_id: int):
    for item in items:
        add_list_item(doc, item, num_id)


def add_numbered(doc: Document, items: Iterable[str], num_id: int):
    for item in items:
        add_list_item(doc, item, num_id)


def add_section(doc: Document, title: str, paragraphs=None, bullets=None, bullet_id=None):
    doc.add_heading(title, level=1)
    for text in paragraphs or []:
        p = doc.add_paragraph(text)
        p.paragraph_format.keep_together = True
    if bullets:
        add_bullets(doc, bullets, bullet_id)


def add_key_files(doc: Document, files: list[tuple[str, str]], bullet_id: int):
    doc.add_heading("关键文件", level=1)
    for path, purpose in files:
        p = add_list_item(doc, f"{path}：{purpose}", bullet_id)
        first = p.runs[0]
        set_run_font(first, latin="Consolas", east_asia="Microsoft YaHei", size=10.5, bold=True, color=COLORS["dark_blue"])


def add_validation(doc: Document, items: list[str], bullet_id: int):
    doc.add_heading("验证要点", level=1)
    add_bullets(doc, items, bullet_id)


VERSION_DOCS = [
    {
        "filename": "01-v1.0-基础测试系统.docx",
        "version": "v1.0",
        "title": "基础测试系统",
        "summary": "建立可完整运行的电影阅历测试流程，让用户从选择场次到获得阅片段位形成闭环。",
        "sections": [
            ("版本概述", ["本版本完成网站的核心产品骨架：用户选择测试题量和电影类型，根据隐藏片名的电影画面完成四选一识别，最终获得百分制成绩与阅片段位。网站不要求登录，打开后即可使用。"], None),
            ("新增功能", None, [
                "测试场次：提供快速 10 部、标准 20 部和深度 30 部三种题量。",
                "电影分类：支持综合、华语、欧美、日韩、动画、科幻、悬疑、恐怖、喜剧与文艺经典。",
                "单阶段答题：选对电影名称即判定答对，不再追加剧情验证题。",
                "即时反馈：作答后立即显示正确状态、电影名称和简短介绍。",
                "结果闭环：统计得分、正确率、最高连胜和分类表现，并生成阅片段位。",
                "本地记录：保存当前测试、历史成绩和最佳成绩，刷新后可继续。",
            ]),
            ("评分方式", ["最终成绩按“正确识别数量 ÷ 总题数”换算为百分制。连续答对用于展示最高连胜，不额外改变基础正确率；不同类型的答题记录用于生成分类能力表现。"], None),
            ("修复与调整", None, [
                "移除“看过／没看过”的主观判断，改为实际片名识别。",
                "删除选对片名后的第二道内容题，降低流程负担。",
                "为返回首页、重新测试和恢复测试提供真实可用的交互。",
                "随机抽题时避免同一场测试内部出现重复电影。",
            ]),
            ("简单技术实现", None, [
                "应用框架：使用 React 19、TypeScript 与 Vite 构建单页应用。",
                "状态结构：QuizSession 保存题目、当前位置、答案、连胜和开始时间。",
                "成绩计算：AnswerRecord 记录每道题的选择与判定，测试结束后统一计算 QuizResult。",
                "浏览器存储：通过 localStorage 保存测试进度与历史记录，无需独立业务后端。",
                "组件拆分：首页、难度选择、答题页、进度条、海报卡片和结果页分别维护。",
            ]),
        ],
        "files": [
            ("src/App.tsx", "控制页面流程、测试状态和结果保存。"),
            ("src/types.ts", "定义电影、题目、答案、成绩和段位类型。"),
            ("src/utils/quiz.ts", "负责抽题、选项随机化和成绩计算。"),
            ("src/data/movies.ts", "提供离线可用的本地示例题库。"),
        ],
        "validation": ["10／20／30 题流程均可完成。", "鼠标、触摸和数字键 1—4 均可答题。", "刷新页面后可恢复未完成测试。", "生产构建能够正常生成。"],
    },
    {
        "filename": "02-v1.1-TMDB实时电影数据库.docx",
        "version": "v1.1",
        "title": "TMDB 实时电影数据库",
        "summary": "由固定本地题库升级为实时电影数据，同时保留离线降级能力和安全的服务端访问方式。",
        "sections": [
            ("版本概述", ["本版本接入 TMDB 电影数据库。每次开始测试时，系统会依据所选分类和难度实时建立候选电影池，并同步片名、上映年份、地区、类型、评分、简介及多张电影图片。"], None),
            ("新增功能", None, [
                "实时题库：测试开始时从 TMDB Discover 接口获取候选电影。",
                "电影资料：展示上映年份、类型、评分、简介和 TMDB 详情地址。",
                "多图来源：同时准备剧照和海报候选，图片失败时自动尝试下一张。",
                "本地降级：接口不可用或凭证缺失时，自动切换到内置题库。",
                "无需用户输入：生产环境由服务端读取 API Key，普通用户直接开始测试。",
            ]),
            ("数据处理", None, [
                "筛选符合分类、票数、评分、热度和年代条件的正式叙事电影。",
                "将 TMDB 类型编号映射为网站使用的中文类型。",
                "根据原始语言归类华语、欧美和日韩电影。",
                "把实时数据转换为统一的 Movie 类型，保证实时题库与本地题库使用同一套界面。",
            ]),
            ("修复内容", None, [
                "解决用户每次打开网站都需要重新输入 API 的问题。",
                "解决本地题库数量有限、分类测试重复度高的问题。",
                "在接口失败时避免测试页面直接中断。",
                "避免将生产密钥写入浏览器构建产物或公开源码。",
            ]),
            ("简单技术实现", None, [
                "前端服务：src/services/tmdb.ts 负责参数生成、响应转换、筛选和抽题。",
                "服务端代理：/api/tmdb/* 由 Netlify Function 转发到 TMDB，并限制允许访问的路径。",
                "请求保护：设置超时、错误状态处理和缓存响应头。",
                "环境变量：生产环境使用 TMDB_API_KEY；本地开发可使用 .env.local。",
                "兼容策略：在线请求失败后调用 createQuiz 生成本地测试。",
            ]),
        ],
        "files": [
            ("src/services/tmdb.ts", "实时电影请求、筛选、数据映射和题目生成。"),
            ("netlify/functions/tmdb.mts", "服务端 TMDB 代理与允许路径校验。"),
            ("src/data/movies.ts", "接口不可用时使用的本地题库。"),
            (".env.example", "本地开发环境变量示例。"),
        ],
        "validation": ["有凭证时使用 TMDB 实时数据。", "无凭证或请求失败时仍能完成本地测试。", "前端构建文件中不包含生产 API Key。", "电影详情字段能够正确展示。"],
    },
    {
        "filename": "03-v1.2-电影化视觉重构.docx",
        "version": "v1.2",
        "title": "电影化视觉重构",
        "summary": "用电影海报、深色胶片质感和克制的科技动效，建立统一而具有辨识度的视觉语言。",
        "sections": [
            ("设计目标", ["视觉方向由普通测试页面升级为“高级电影质感融合科技感”。整体使用深海军蓝、半透明玻璃层、电影海报和编辑式大标题，减少廉价霓虹效果，优先保证文字与答题区域清晰。"], None),
            ("首页更新", None, [
                "海报背景：多行电影海报持续流动，构成电影资料库般的空间感。",
                "电影化标题：使用大字号衬线字体与分层灰白色彩形成主视觉。",
                "玻璃导航：导航、模式选择和主要按钮采用轻量液态玻璃效果。",
                "入场动效：标题、简介和按钮依次上升淡入。",
                "信息避让：海报仅作为背景，不遮挡核心文案和操作区域。",
            ]),
            ("答题页更新", None, [
                "左侧完整展示电影画面，右侧集中呈现题目、选项和反馈。",
                "背景使用当前电影与其他候选海报形成低干扰动态氛围。",
                "选项具有 hover、focus、disabled、正确和错误状态。",
                "顶部显示题号、进度、返回首页、重新测试和语言切换。",
                "移动端重新排列内容，优先保证题目和选项可读。",
            ]),
            ("修复内容", None, [
                "修复海报容器裁切过多、无法看到主要画面的排版问题。",
                "修复答题背景过暗、用户无法感知电影元素的问题。",
                "修复中文首页大标题标点与下一行文字重叠。",
                "统一首页和答题页的字体、边框、阴影与动效节奏。",
            ]),
            ("简单技术实现", None, [
                "使用 CSS 自定义属性维护颜色、字体、玻璃透明度和动画参数。",
                "使用 CSS keyframes 实现海报轨道和 fade-rise 入场动画。",
                "根据 html 的 lang 属性分别调整中英文标题行距。",
                "通过 clamp、媒体查询和弹性布局适配桌面与手机。",
                "遵循 prefers-reduced-motion，在用户减少动态效果时降低动画强度。",
            ]),
        ],
        "files": [
            ("src/styles.css", "视觉变量、页面布局、响应式规则和动画。"),
            ("src/components/HomeScreen.tsx", "首页海报墙、导航和测试场次入口。"),
            ("src/components/QuizScreen.tsx", "电影化答题布局与反馈区域。"),
            ("src/components/PosterCard.tsx", "电影画面与隐藏片名信息卡。"),
        ],
        "validation": ["桌面和手机宽度下文字均不重叠。", "键盘焦点状态清晰可见。", "背景动效不遮挡选项。", "中文与英文标题均保持合理换行。"],
    },
    {
        "filename": "04-v1.3-题库性能与图片稳定性.docx",
        "version": "v1.3",
        "title": "题库性能与图片稳定性",
        "summary": "缩短首页和题库等待时间，并通过多级图片回退保证电影画面不会留下空白。",
        "sections": [
            ("版本概述", ["本版本集中处理两个高频问题：首页海报加载缓慢，以及进入测试时题库等待时间过长。优化后，页面会先使用可立即显示的资源，再在后台更新实时数据。"], None),
            ("加载性能优化", None, [
                "首屏优先：先使用本地或浏览器缓存的海报池渲染首页。",
                "后台刷新：页面稳定后再异步获取新的 TMDB 海报，供下次访问使用。",
                "请求并发：多个候选页并行获取，减少顺序等待时间。",
                "缓存复用：一定时间内复用首页海报数据，避免重复请求。",
                "按需补图：进入测试后在后台为缺少无字剧照的电影补充图片。",
            ]),
            ("图片可靠性", None, [
                "每部电影维护多张剧照与海报候选地址。",
                "图片超时或加载失败后自动切换到下一候选。",
                "首页海报还会借用海报池中的其他图片作为紧急回退。",
                "所有候选均失败时显示尺寸稳定的占位视觉，防止布局坍塌。",
                "答题图优先使用无语言标记的剧照，降低片名泄露概率。",
            ]),
            ("修复内容", None, [
                "修复首页加载期间出现大片空白卡位。",
                "修复日韩、悬疑等分类偶发没有海报。",
                "修复单一图片地址失效后整题无法显示。",
                "减少重复进入测试时再次等待完整题库请求。",
            ]),
            ("简单技术实现", None, [
                "ResilientPosterImage 按顺序尝试 sources 数组，并支持加载超时。",
                "homePosterCandidates 与 moviePosterCandidates 统一生成候选图片列表。",
                "localStorage 保存首页海报池及时间戳，超过有效期后后台刷新。",
                "Promise.allSettled 允许部分 TMDB 请求失败而不影响其他结果。",
                "图片 URL 使用不同 TMDB 尺寸作为同一图片的多级回退。",
            ]),
        ],
        "files": [
            ("src/components/ResilientPosterImage.tsx", "多候选图片加载与错误回退。"),
            ("src/utils/posters.ts", "电影图片候选和海报池工具。"),
            ("src/components/HomeScreen.tsx", "首屏缓存海报与后台刷新。"),
            ("src/services/tmdb.ts", "并发题库请求和剧照补充。"),
        ],
        "validation": ["断开某张图片后会自动显示备用图。", "首页首次渲染不依赖实时接口完成。", "不同分类均能生成完整测试。", "所有图片失败时布局仍保持完整。"],
    },
    {
        "filename": "05-v1.4-难度系统与选片扩展.docx",
        "version": "v1.4",
        "title": "难度系统与选片扩展",
        "summary": "让题量、电影类型和玩家身份共同决定题目难度，扩大电影史覆盖并降低重复率。",
        "sections": [
            ("版本概述", ["用户选择 10、20 或 30 部电影后，会先进入身份选择，再根据身份与类型实时生成题目。难度不再单纯依赖冷门电影，而是综合作品知名度、观众样本、年代、地区和干扰项相似度。"], None),
            ("三种玩家身份", None, [
                "入门菜鸟：以高知名度、观众基础较大的电影为主，适合轻松热身。",
                "略知一二：扩大年代、地区和类型范围，并提高干扰项相似度。",
                "阅片无数：覆盖更多电影传统、早期经典和作者电影，但仍保留基本口碑与观众样本门槛。",
            ]),
            ("选片范围", None, [
                "覆盖整个电影发展历史，不以“最近上映”作为主要标准。",
                "综合模式平衡华语、欧美、日韩以及不同类型，避免单一区域占满整套题。",
                "排除演唱会、巡演、音乐现场、电视电影和多数纪录片等非目标内容。",
                "不为了提高难度而大量选择低分、极冷门或几乎无人看过的作品。",
                "本地维护代表性作品和人工排除名单，用于校正自动筛选的边界情况。",
            ]),
            ("干扰项改进", None, [
                "优先选择相近地区、原始语言、年代与类型的电影。",
                "比较票数和热度，避免一个全球名作搭配三个明显无关的小众选项。",
                "考虑片名长度，减少仅凭标题形式即可排除答案。",
                "四个选项保持固定身份，语言切换不会改变顺序或正确答案。",
            ]),
            ("防重复机制", None, [
                "浏览器记录曾经出现过的 TMDB 电影 ID，新测试优先选择未出现的电影。",
                "在同一场测试中对电影 ID 去重。",
                "候选不足时分层放宽条件，而不是直接重复最近题目。",
                "记录不同身份下的答题表现，为后续难度校准积累本机样本。",
            ]),
            ("简单技术实现", None, [
                "使用不同难度档案控制评分、票数、热度和可抽取页数。",
                "按年代权重从早期电影、黄金年代、现代电影和近年作品中组合题单。",
                "通过 distractorScore 对类型重叠、语言、年份、票数、热度和标题长度加权。",
                "使用 movieCuration 维护人工精选与排除 TMDB ID。",
                "使用 performance 工具保存已出现电影和本地答题统计。",
            ]),
        ],
        "files": [
            ("src/components/DifficultyScreen.tsx", "身份与电影类型选择界面。"),
            ("src/services/tmdb.ts", "难度档案、年代配额、筛选和干扰项评分。"),
            ("src/data/movieCuration.ts", "精选作品和人工排除名单。"),
            ("src/utils/performance.ts", "出现记录与分身份答题表现。"),
        ],
        "validation": ["选择题量后先进入身份选择。", "三种身份生成的电影范围明显不同。", "题单覆盖多个年代和地区。", "连续多次测试的重复率明显下降。"],
    },
    {
        "filename": "06-v1.5-答题反馈与结果体验.docx",
        "version": "v1.5",
        "title": "答题反馈与结果体验",
        "summary": "让用户在每道题后获得足够但不过量的电影信息，并在测试结束后清楚理解自己的表现。",
        "sections": [
            ("版本概述", ["本版本加强作答后的信息反馈和结果页表达。用户不仅能看到对错，还能立即了解电影年份、地区、类型、评分和简介，并通过 TMDB 链接继续查看完整资料。"], None),
            ("单题反馈", None, [
                "正确状态：用清晰的颜色、图标和文字说明是否识别成功。",
                "正确片名：答错时直接显示正确答案。",
                "简短资料：展示电影简介、年份、地区、类型和 TMDB 评分。",
                "电影链接：在详细信息下方提供明显的 TMDB 官方资料入口。",
                "继续流程：反馈区域末尾提供“下一部电影”或“查看我的段位”。",
            ]),
            ("结果页面", None, [
                "显示百分制总分、总体正确率和最高连胜。",
                "根据得分和所选身份授予对应阅片段位。",
                "用分类数据和雷达图展示不同电影类型的表现。",
                "显示正确识别数量、印象模糊数量和测试来源。",
                "提供再测一次、换类型和分享结果入口。",
            ]),
            ("数据与隐私", None, [
                "当前测试、历史成绩和最佳成绩保存在浏览器 localStorage。",
                "网站无需注册，不主动上传用户的答题记录。",
                "结果页不公开具体题目答案，分享时避免泄露题库。",
                "清除浏览器站点数据后，本地历史记录也会被删除。",
            ]),
            ("修复内容", None, [
                "修复电影详情链接不明显、用户难以发现的问题。",
                "将 TMDB 链接放到电影详细信息下方，并使用完整可点击区域。",
                "调整反馈信息层级，避免简介、标签和下一题按钮挤在一起。",
                "不再使用缺少真实样本支持的“超越多少玩家”模拟百分位。",
            ]),
            ("简单技术实现", None, [
                "AnswerRecord 记录选择、正确性、类型和得分。",
                "calculateResult 汇总总分、正确率、连胜和分类成绩。",
                "ResultScreen 根据 Rank 阈值显示段位卡与评价。",
                "RadarChart 将分类成绩转化为可视化能力轮廓。",
                "TMDB 详情 URL 使用电影 tmdbId 生成，并随语言切换 locale。",
            ]),
        ],
        "files": [
            ("src/components/QuizScreen.tsx", "单题反馈、电影资料和 TMDB 链接。"),
            ("src/components/ResultScreen.tsx", "测试结果、段位和历史表现。"),
            ("src/components/RadarChart.tsx", "分类能力雷达图。"),
            ("src/utils/quiz.ts", "答案记录和结果计算。"),
        ],
        "validation": ["答题后能够打开对应电影的 TMDB 页面。", "结果统计与实际答题数量一致。", "分享内容不包含题目答案。", "历史成绩关闭页面后仍可读取。"],
    },
    {
        "filename": "07-v1.6-Netlify部署与数据安全.docx",
        "version": "v1.6",
        "title": "Netlify 部署与数据安全",
        "summary": "把本地应用发布为可通过链接直接访问的网站，并建立 GitHub 到 Netlify 的持续更新流程。",
        "sections": [
            ("部署目标", ["项目通过 GitHub 管理源码，并由 Netlify 构建和发布。代码推送到 main 分支后，Netlify 可以自动执行生产构建并更新在线网站。部署说明不区分地区版本。"], None),
            ("部署流程", None, [
                "将本地项目初始化为 Git 仓库并推送到 GitHub。",
                "在 Netlify 创建项目并连接对应 GitHub 仓库。",
                "生产分支使用 main，构建设置由 netlify.toml 管理。",
                "Netlify 执行 npm run build:netlify，并发布 release 目录。",
                "后续代码更新经过提交和推送后触发自动重新部署。",
            ]),
            ("构建与路由", None, [
                "Node.js 版本固定为 22，减少不同构建环境造成的差异。",
                "Vite 生产资源输出到 release，并使用带哈希的静态文件名。",
                "单页应用路由统一回退到 index.html，刷新页面不会返回 404。",
                "静态 assets 使用长期缓存，HTML 与接口保持适合更新的缓存策略。",
            ]),
            ("TMDB 密钥安全", None, [
                "生产密钥保存为 Netlify 环境变量 TMDB_API_KEY。",
                "前端只访问 /api/tmdb/*，不直接读取生产密钥。",
                "Netlify Function 在服务端附加 API Key，并限制可访问的 TMDB 路径。",
                "环境变量不要使用 VITE_ 前缀，否则可能被打包进浏览器代码。",
                ".env.local 仅用于本地开发，并由 .gitignore 排除。",
            ]),
            ("常见问题", None, [
                "GitHub TLS 超时：检查代理是否允许 PowerShell 和 GitHub CLI 访问 443 端口。",
                "更新未生效：确认本地提交已推送、Netlify 最新部署成功，并刷新浏览器缓存。",
                "环境变量未生效：保存变量后重新触发生产部署。",
                "本地 127.0.0.1 拒绝连接：先执行 npm run dev 或 npm run preview。",
                "npm 被系统要求选择应用打开：在 PowerShell 中直接调用有效的 npm.cmd。",
            ]),
            ("简单技术实现", None, [
                "netlify.toml 声明构建、发布目录、Functions 目录、重定向和安全响应头。",
                "tmdb.mts 使用 Netlify Functions v2 接收 GET 请求并转发允许的 TMDB 路径。",
                "接口设置 10 秒超时，并对成功响应提供 CDN 缓存与 stale-while-revalidate。",
                "同步脚本封装 git add、commit 和 push，减少重复输入命令。",
            ]),
        ],
        "files": [
            ("netlify.toml", "构建、发布、路由、缓存和安全头配置。"),
            ("netlify/functions/tmdb.mts", "安全的服务端 TMDB 代理。"),
            ("package.json", "本地与 Netlify 构建脚本。"),
            ("同步到GitHub.cmd", "本地代码提交与推送辅助脚本。"),
        ],
        "validation": ["Netlify 构建命令成功完成。", "公开链接可正常打开并刷新。", "TMDB 实时题库在部署环境中可用。", "浏览器构建产物中不包含生产密钥。"],
    },
    {
        "filename": "08-v1.7-中英文切换与体验修复.docx",
        "version": "v1.7",
        "title": "中英文切换与体验修复",
        "summary": "实现中文与英文自由切换，并确保答题途中切换语言不会改变选项身份、判分或当前进度。",
        "sections": [
            ("版本概述", ["网站新增完整的中英文界面。语言选择会保存在浏览器中，也可以通过 ?lang=zh 或 ?lang=en 直接打开指定语言。测试开始后，用户仍可随时切换语言。"], None),
            ("新增功能", None, [
                "首页、难度选择、加载页、答题页和结果页提供中英文文案。",
                "电影类型、玩家身份、难度、地区和类型标签均可翻译。",
                "TMDB 同时获取中文与英文片名、简介和干扰项。",
                "电影详情链接会根据当前语言打开对应的 TMDB 页面。",
                "语言偏好保存在 localStorage，重新访问时自动恢复。",
            ]),
            ("答题途中切换", None, [
                "每个选项使用稳定 ID，而不是使用当前显示文字作为身份。",
                "切换语言只更新片名文字，不重新随机排列四个选项。",
                "已经选择的选项、正确答案和得分不会改变。",
                "电影简介和答题反馈随当前语言同步更新。",
                "旧版未完成测试缺少双语数据时重新生成，历史成绩继续保留。",
            ]),
            ("修复内容", None, [
                "修复题目变成中文但选项仍保持英文。",
                "修复中文首页标题中逗号与下一行“阅”字重叠。",
                "修复 TMDB 偶发返回空中文标题导致选项无文字。",
                "为视觉上容易误认为空白的《一一》显示“一一（Yi Yi）”。",
                "空翻译会依次回退到原始片名、另一语言片名或安全占位文字。",
            ]),
            ("简单技术实现", None, [
                "LanguageProvider 使用 React Context 管理语言，并同步 html lang 与页面标题。",
                "Movie 增加 localizedTitles、localizedDistractors 和 localizedSynopses。",
                "TMDB 候选页分别请求 zh-CN 与 en-US，再按电影 ID 合并双语资料。",
                "QuizScreen 使用 answer 和 distractor-N 作为稳定选项键。",
                "CSS 根据 html[lang] 为中文标题设置独立行距和换行结构。",
                "firstNonBlank 对所有片名进行 trim 与非空回退。",
            ]),
        ],
        "files": [
            ("src/i18n.tsx", "语言状态、标签翻译和偏好保存。"),
            ("src/components/LanguageSwitch.tsx", "中文／英文切换控件。"),
            ("src/components/QuizScreen.tsx", "稳定选项 ID 与双语显示。"),
            ("src/services/tmdb.ts", "中英文电影资料合并与空标题回退。"),
            ("src/styles.css", "中英文排版差异与标题防重叠。"),
        ],
        "validation": ["答题前后均可切换语言。", "选项顺序与已选状态保持不变。", "中文标题无重叠。", "所有四个选项始终具有可读片名。", "类型检查与生产构建通过。"],
    },
    {
        "filename": "09-v1.8-好友挑战与个人阅片档案.docx",
        "version": "v1.8",
        "title": "好友挑战与个人阅片档案",
        "summary": "增加同题好友挑战、长期阅片能力画像和退出确认，让测试更适合分享，也更值得持续回来使用。",
        "sections": [
            ("版本概述", ["本版本将一次性的电影测试扩展为可传播、可积累的长期体验。用户完成测试后可以邀请好友回答完全相同的电影；历史答题会逐渐形成个人阅片档案；答题途中离开时会先确认如何处理当前进度。"], None),
            ("好友同题挑战", None, [
                "测试结束后可生成专属挑战链接，链接包含同一组电影题目和发起者成绩。",
                "朋友打开链接后先看到挑战题量、身份、类型和需要挑战的正确数量。",
                "朋友完成后比较双方正确数量、共同答对、只有一方答对和擅长类型。",
                "挑战链接无需账号和业务数据库，不包含发起者的历史成绩或其他个人信息。",
                "挑战完成者还可以使用自己的成绩再次生成新的同题挑战。",
            ]),
            ("个人阅片档案", None, [
                "统计擅长地区、擅长年代、擅长类型和最容易失误的电影类型。",
                "统计已识别的不同电影数量、累计作答次数和历史正确率。",
                "根据历史测试记录显示曾完成的最高身份难度。",
                "用折线趋势展示最近 12 场测试的阅片能力变化。",
                "档案数据继续保存在当前浏览器，不要求注册，也不会主动上传。",
            ]),
            ("退出确认", None, [
                "答题中点击品牌、返回首页、重新测试或按 Esc 时不再立即离开。",
                "选择“保存并返回”会保留已完成题目，之后可从首页继续。",
                "选择“放弃本次测试”会删除未完成场次并返回首页。",
                "重新测试时可选择放弃当前进度并立即开始同配置的新场次。",
                "确认窗口始终提供“继续答题”，避免鼠标或触摸误操作。",
            ]),
            ("简单技术实现", None, [
                "挑战载荷使用浏览器 CompressionStream 压缩为 gzip，并写入 URL fragment；片段不会发送给 Netlify。",
                "ChallengePayload 保存电影、设置和发起者答对电影 ID，ChallengeComparison 在好友完成后按 ID 比较。",
                "MoviePerformanceStats 增加年份、地区和类型，用于长期聚合个人能力画像。",
                "getViewingProfile 从本地电影表现和 QuizResult 历史中生成地区、年代、类型和趋势指标。",
                "QuizSession 增加可选 challenge 上下文，未完成挑战仍可通过 localStorage 恢复。",
                "退出确认使用受控模态状态区分返回、放弃和重新测试。",
            ]),
            ("兼容性说明", ["旧历史记录仍可正常显示分数与趋势。早期电影表现记录没有年份、地区字段，因此地区和年代画像会在用户继续答题后逐渐补全。挑战链接依赖现代浏览器的压缩能力；不支持压缩时会自动使用未压缩的兼容格式。"], None),
        ],
        "files": [
            ("src/utils/challenge.ts", "挑战载荷压缩、链接解析、挑战会话和结果比较。"),
            ("src/components/ChallengeIntro.tsx", "好友挑战介绍与接受挑战页面。"),
            ("src/components/ProfileScreen.tsx", "个人阅片档案和能力趋势页面。"),
            ("src/utils/performance.ts", "地区、年代、类型与历史难度聚合。"),
            ("src/components/QuizScreen.tsx", "保存、放弃和重新测试确认窗口。"),
            ("src/components/ResultScreen.tsx", "挑战链接生成和双方结果对比。"),
        ],
        "validation": ["挑战链接能够还原相同电影题目。", "朋友完成后四类对比数据正确。", "档案指标随新答题记录更新。", "返回首页保留进度，放弃测试清除进度。", "中英文界面和手机布局均保持可用。", "类型检查与生产构建通过。"],
    },
]


def build_version_doc(spec: dict):
    doc = Document()
    configure_document(doc, f"{spec['version']} · {spec['title']}")
    doc.core_properties.title = f"{spec['version']} {spec['title']}"
    bullet_id = create_numbering(doc, bullet=True)
    add_title_block(doc, spec["version"], spec["title"], spec["summary"])
    for title, paragraphs, bullets in spec["sections"]:
        add_section(doc, title, paragraphs=paragraphs, bullets=bullets, bullet_id=bullet_id)
    add_key_files(doc, spec["files"], bullet_id)
    add_validation(doc, spec["validation"], bullet_id)
    path = OUTPUT_DIR / spec["filename"]
    doc.save(path)
    return path


def build_index():
    doc = Document()
    configure_document(doc, "版本文档索引")
    doc.core_properties.title = "光影鉴赏局版本文档索引"
    bullet_id = create_numbering(doc, bullet=True)
    number_id = create_numbering(doc, bullet=False)
    add_title_block(
        doc,
        "DOCUMENTATION",
        "光影鉴赏局版本文档索引",
        "汇总电影阅历测试网站从核心流程、实时题库到视觉、性能、难度、部署与双语体验的主要迭代。",
        status="持续维护",
    )

    add_section(doc, "文档说明", paragraphs=[
        "本目录面向普通用户、产品维护者和后续开发者。每份文档只概括该阶段最重要的新增功能、修复内容与技术实现，不替代源代码或接口文档。",
    ], bullets=[
        "用户读者：重点阅读“新增功能”“修复内容”和“用户可见变化”。",
        "维护者：重点阅读“简单技术实现”“关键文件”和“验证要点”。",
        "部署说明：Netlify 仅作为网站部署流程介绍，不区分地区版本。",
    ], bullet_id=bullet_id)

    doc.add_heading("版本总览", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["版本", "主题", "主要更新", "核心技术"]
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, COLORS["light_fill"])
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(text)
        set_run_font(run, size=9.5, bold=True, color=COLORS["navy"])
    mark_repeat_table_header(table.rows[0])

    overview_rows = [
        ("v1.0", "基础测试系统", "完整答题、评分、段位和本地记录", "React 状态与 localStorage"),
        ("v1.1", "TMDB 实时数据库", "实时电影资料、多图和本地降级", "TMDB API 与服务端代理"),
        ("v1.2", "电影化视觉重构", "海报墙、玻璃界面和沉浸式答题", "响应式 CSS 与动画"),
        ("v1.3", "性能与图片稳定", "缓存、并发加载和多级图片回退", "异步请求与候选队列"),
        ("v1.4", "难度与选片扩展", "身份难度、电影史覆盖和防重复", "加权筛选与干扰项评分"),
        ("v1.5", "反馈与结果体验", "电影详情、TMDB 链接和段位结果", "答题记录与结果聚合"),
        ("v1.6", "Netlify 部署", "持续部署、路由、缓存和密钥安全", "GitHub、Netlify Functions"),
        ("v1.7", "中英文与体验修复", "双语切换、稳定选项和排版修复", "React Context 与双语数据"),
        ("v1.8", "挑战与个人档案", "同题好友挑战、能力画像和退出确认", "压缩链接与本地数据聚合"),
    ]
    for row_data in overview_rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row_data):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            if idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            set_run_font(run, size=9.2, bold=idx == 0, color=COLORS["body"])
    set_table_geometry(table, [900, 1770, 3510, 3180])

    doc.add_heading("建议阅读顺序", level=1)
    add_numbered(doc, [
        "第一次了解项目：先阅读 v1.0、v1.4 和 v1.5，掌握用户实际经历的完整流程。",
        "准备维护题库：阅读 v1.1、v1.3 和 v1.4，了解实时数据、图片与筛选机制。",
        "准备修改界面：阅读 v1.2 和 v1.7，了解视觉系统、响应式布局与双语状态。",
        "准备更新网站：阅读 v1.6，完成构建、推送、环境变量与部署验证。",
    ], number_id)

    doc.add_heading("当前能力摘要", level=1)
    add_bullets(doc, [
        "支持 10、20、30 部电影三种题量和十类电影范围。",
        "支持入门菜鸟、略知一二、阅片无数三种难度身份。",
        "支持 TMDB 实时数据、本地题库降级和多图片回退。",
        "支持中英文自由切换、键盘答题、历史记录和段位结果。",
        "支持好友同题挑战、双方结果比较和长期个人阅片档案。",
        "支持 GitHub 管理与 Netlify 自动构建发布。",
    ], bullet_id)

    path = OUTPUT_DIR / "00-版本文档索引.docx"
    doc.save(path)
    return path


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    for existing in OUTPUT_DIR.glob("*.docx"):
        existing.unlink()
    paths = [build_index()]
    paths.extend(build_version_doc(spec) for spec in VERSION_DOCS)
    for path in paths:
        print(path)


if __name__ == "__main__":
    main()
